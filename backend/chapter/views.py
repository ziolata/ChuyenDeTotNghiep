from django.shortcuts import render
from rest_framework import generics, permissions, status, mixins
from rest_framework.permissions import AllowAny, IsAuthenticated, IsAdminUser
from rest_framework.response import Response
from .models import *
from .serializers import ChapterSerializer, ChapterDetailSerializer, ChapterContentSerializer
from rest_framework.pagination import PageNumberPagination
from PyPDF2 import PdfReader
from docx import Document
# Create your views here.
class StandardResultsSetPagination(PageNumberPagination):
    page_size = 8
    page_size_query_param = 'page_size'
    max_page_size = 1000
    
from docx import Document
from PyPDF2 import PdfReader

class ChapterDetailView(generics.RetrieveAPIView):
    queryset = Chapter.objects.all()
    serializer_class = ChapterDetailSerializer
    permission_classes = [permissions.AllowAny]

    def read_file(self, instance):
        try:
            if instance.content.name.endswith('.docx'):
                with instance.content.open(mode='rb') as file:
                    doc = Document(file)
                    paragraphs = []
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            paragraph_text = f"<p>{paragraph.text}</p>"
                            paragraphs.append(paragraph_text)
                    file_content = '\n'.join(paragraphs)
                    return file_content
            elif instance.content.name.endswith('.pdf'):
                with instance.content.open(mode='rb') as file:
                    reader = PdfReader(file)
                    text = ""
                    for page_num in range(len(reader.pages)):
                        page = reader.pages[page_num]
                        text += "<p>" + page.extract_text().replace("\n", "</p><p>") + "</p>"
                    return text
            elif instance.content.name.endswith('.txt'):
                with instance.content.open(mode='r') as file:
                    text = file.read()
                    paragraphs = ["<p>" + line.strip() + "</p>" for line in text.split("\n") if line.strip()]
                    file_content = '\n'.join(paragraphs)
                    return file_content
            else:
                return "Unsupported file format"
        except Exception as e:
            return str(e)

    def get_next_chapter_id(self, current_chapter):
        try:
            novel_chapters = Chapter.objects.filter(novel=current_chapter.novel)
            novel_chapters = novel_chapters.order_by('number')
            current_chapter_index = list(novel_chapters).index(current_chapter)
            next_chapter = novel_chapters[current_chapter_index + 1]
            return next_chapter.id
        except IndexError:
            return None

    def get_previous_chapter_id(self, current_chapter):
        try:
            novel_chapters = Chapter.objects.filter(novel=current_chapter.novel)
            novel_chapters = novel_chapters.order_by('number')
            current_chapter_index = list(novel_chapters).index(current_chapter)
            if current_chapter_index == 0:
                return None 
            else:
                previous_chapter = novel_chapters[current_chapter_index - 1]
                return previous_chapter.id
        except IndexError:
            return None

    def retrieve(self, request, *args, **kwargs):
        try:
            instance = self.get_object()
            file_content = self.read_file(instance)

            serializer = self.get_serializer(instance)
            data = serializer.data
            data['content'] = file_content

            next_chapter_id = self.get_next_chapter_id(instance)
            previous_chapter_id = self.get_previous_chapter_id(instance)

            data['next_chapter_id'] = next_chapter_id
            data['previous_chapter_id'] = previous_chapter_id

            return Response(data)
        except Chapter.DoesNotExist:
            return Response({'details': f"Chapter with id {kwargs['pk']} not found"}, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            return Response({'details': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)



   
    

class ChapterView(generics.ListAPIView):
    queryset = Chapter.objects.all().order_by('createdAt')
    serializer_class = ChapterSerializer
    permission_classes = [permissions.AllowAny]
    pagination_class = StandardResultsSetPagination
    def list(self, request, pk=None):
        try:
            chapters = Chapter.objects.filter(novel__pk=pk).order_by('createdAt')
            
            # Phân trang cho queryset
            page = self.paginate_queryset(chapters)
            if page is not None:
                serializer = ChapterSerializer(page, many=True)
                return self.get_paginated_response(serializer.data)

            serializer = ChapterSerializer(chapters, many=True)
            return Response(serializer.data)
        except Exception as e:
            return Response({'details': f"{e}"}, status=status.HTTP_204_NO_CONTENT)


class ChapterContentView(generics.RetrieveAPIView):
    queryset = ChapterContent.objects.all()
    serializer_class = ChapterContentSerializer
    permission_classes = [permissions.AllowAny]

    def get(self, request, *args, **kwargs):
        instance = self.get_object()
        try:
            with instance.content.open(mode='rb') as file:
                doc = Document(file)
                paragraphs = []
                for paragraph in doc.paragraphs:
                    paragraphs.append(paragraph._element.xml)
                file_content = '\n'.join(paragraphs)
        except Exception as e:
            return Response({'error': str(e)}, status=500)
        return Response({'content': file_content})
    


